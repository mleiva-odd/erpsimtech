from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "docs" / "manuales" / "Manual_de_Primer_Uso_SIMTECH.docx"
LOGO_PATH = ROOT / "public" / "logo.png"
POS_IMAGE = ROOT / "public" / "pos.jpg"
INVENTORY_IMAGE = ROOT / "public" / "inventory.jpg"
ANALYTICS_IMAGE = ROOT / "public" / "analysis.jpg"

BLUE = "1F5FAF"
BLUE_DARK = "163A63"
BLUE_SOFT = "EAF2FF"
GREEN_SOFT = "EAF7EE"
AMBER_SOFT = "FFF5DD"
GRAY_SOFT = "F4F7FA"
TEXT = RGBColor(40, 52, 66)
MUTED = RGBColor(99, 115, 129)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def hide_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "nil")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT

    for style_name, size, color in [
        ("Title", 24, BLUE_DARK),
        ("Heading 1", 16, BLUE_DARK),
        ("Heading 2", 12, BLUE_DARK),
        ("Heading 3", 10.5, BLUE_DARK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.font.color.rgb = TEXT

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("SIMTECH | Manual de primer uso para clientes")
    run.font.name = "Aptos"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED


def add_spacer(doc: Document, size: float = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(size)


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.color.rgb = TEXT
        rest = p.add_run(text[len(bold_prefix):])
        rest.font.color.rgb = TEXT
    else:
        run = p.add_run(text)
        run.font.color.rgb = TEXT


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if ":" in item:
            prefix, rest = item.split(":", 1)
            r1 = p.add_run(prefix + ":")
            r1.bold = True
            r2 = p.add_run(rest)
            r2.font.color.rgb = TEXT
        else:
            p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if ":" in item:
            prefix, rest = item.split(":", 1)
            r1 = p.add_run(prefix + ":")
            r1.bold = True
            r2 = p.add_run(rest)
            r2.font.color.rgb = TEXT
        else:
            p.add_run(item)


def add_banner_heading(doc: Document, title: str, subtitle: str | None = None) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(7.0)
    hide_table_borders(table)

    cell = table.cell(0, 0)
    set_cell_shading(cell, BLUE)
    set_cell_margins(cell, top=100, start=160, bottom=110, end=160)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(255, 255, 255)

    if subtitle:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(9.5)
        run2.font.color.rgb = RGBColor(228, 238, 249)

    add_spacer(doc, 6)


def add_callout(doc: Document, title: str, lines: list[str], fill: str = BLUE_SOFT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    hide_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=110, start=140, bottom=100, end=140)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    head = cell.paragraphs[0]
    head.paragraph_format.space_after = Pt(4)
    r1 = head.add_run(title)
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor.from_string(BLUE_DARK)

    for idx, line in enumerate(lines):
        p = cell.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Pt(14)
        p.paragraph_format.space_after = Pt(2 if idx < len(lines) - 1 else 0)
        p.add_run(line)
    add_spacer(doc, 6)


def add_image(doc: Document, path: Path, width: float, caption: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    p.paragraph_format.space_after = Pt(4)

    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        cr = cp.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(8.5)
        cr.font.color.rgb = MUTED


def add_two_col_table(doc: Document, headers: tuple[str, str], rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(2.1)
    table.columns[1].width = Inches(4.9)

    header_cells = table.rows[0].cells
    for idx, value in enumerate(headers):
        cell = header_cells[idx]
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(value)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    for left, right in rows:
        row = table.add_row().cells
        values = [left, right]
        for idx, value in enumerate(values):
            cell = row[idx]
            set_cell_margins(cell)
            if idx == 0:
                set_cell_shading(cell, GRAY_SOFT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.size = Pt(10)

    add_spacer(doc, 6)


def add_checklist_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.7)
    table.columns[1].width = Inches(3.0)
    table.columns[2].width = Inches(2.3)

    headers = ["Área", "Qué revisar", "Resultado esperado"]
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(value)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)

    for area, check, outcome in rows:
        row = table.add_row().cells
        for idx, value in enumerate([area, check, outcome]):
            cell = row[idx]
            set_cell_margins(cell, top=90, start=110, bottom=90, end=110)
            if idx == 0:
                set_cell_shading(cell, GRAY_SOFT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if idx == 0:
                run = p.add_run(value)
                run.bold = True
            else:
                p.add_run(value)

    add_spacer(doc, 6)


def add_cover(doc: Document) -> None:
    if LOGO_PATH.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(str(LOGO_PATH), width=Inches(1.1))
        p_logo.paragraph_format.space_after = Pt(8)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(3)
    title = p_title.add_run("Manual de primer uso")
    title.bold = True
    title.font.size = Pt(24)
    title.font.color.rgb = RGBColor.from_string(BLUE_DARK)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(10)
    sub = p_sub.add_run("SIMTECH ERP/POS | Guía práctica para clientes")
    sub.font.size = Pt(12)
    sub.font.color.rgb = MUTED

    summary = doc.add_table(rows=1, cols=1)
    hide_table_borders(summary)
    cell = summary.cell(0, 0)
    set_cell_shading(cell, BLUE_SOFT)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(
        "Este documento acompaña a los usuarios en su configuración inicial, la carga del catálogo base y el inicio de la operación diaria."
    )
    r.font.size = Pt(10.5)
    r.font.color.rgb = TEXT

    if POS_IMAGE.exists():
        add_spacer(doc, 4)
        add_image(doc, POS_IMAGE, 6.55, "Vista ilustrativa del módulo de punto de venta.")

    meta = doc.add_table(rows=2, cols=2)
    hide_table_borders(meta)
    meta.autofit = False
    meta.columns[0].width = Inches(1.55)
    meta.columns[1].width = Inches(5.15)
    values = [
        ("Dirigido a", "Administradores, supervisores, cajeros y personal operativo del cliente."),
        ("Versión", f"Edición generada el {date.today().strftime('%d/%m/%Y')}."),
    ]
    for row_idx, (left, right) in enumerate(values):
        left_cell = meta.cell(row_idx, 0)
        right_cell = meta.cell(row_idx, 1)
        set_cell_shading(left_cell, GRAY_SOFT)
        set_cell_margins(left_cell)
        set_cell_margins(right_cell)
        lp = left_cell.paragraphs[0]
        lr = lp.add_run(left)
        lr.bold = True
        rp = right_cell.paragraphs[0]
        rp.add_run(right)

    doc.add_page_break()


def build_manual() -> None:
    doc = Document()
    style_document(doc)
    add_cover(doc)

    add_banner_heading(doc, "1. Cómo leer este manual", "Orden sugerido para poner en marcha el sistema sin omitir pasos clave.")
    add_body(
        doc,
        "Este manual está pensado para el primer contacto del cliente con SIMTECH. La idea es que el negocio quede listo para operar desde el primer día y que cada usuario entienda en qué módulo trabajar.",
    )
    add_callout(
        doc,
        "Ruta recomendada del arranque",
        [
            "Registrar o recibir las credenciales de acceso.",
            "Configurar datos del negocio, sucursales, usuarios y permisos.",
            "Cargar productos, clientes y proveedores.",
            "Definir bancos y medios de cobro si se usarán tarjeta o transferencia.",
            "Abrir caja, hacer una venta de prueba y confirmar reportes.",
        ],
        fill=GREEN_SOFT,
    )
    add_checklist_table(
        doc,
        [
            ("Acceso", "Correo y contraseña del administrador principal.", "El usuario puede ingresar a /login y ver el lanzador de aplicaciones."),
            ("Datos fiscales", "Nombre comercial, NIT, teléfono, dirección y mensaje de ticket.", "El negocio queda listo para identificarse en tickets y configuraciones."),
            ("Operación", "Listado de productos, clientes base y proveedores principales.", "El equipo puede empezar a vender y recibir inventario."),
            ("Cobros", "Métodos de pago permitidos y cuentas bancarias disponibles.", "POS puede cobrar sin bloquear ventas por datos incompletos."),
        ],
    )

    add_banner_heading(doc, "2. Primer acceso al sistema", "Lo que debe hacer el cliente en su primera sesión.")
    add_numbered(
        doc,
        [
            "Ingreso a la plataforma: si la empresa es nueva, el alta inicial solicita datos de empresa, cuenta administradora y primera sucursal. Si la empresa ya fue creada por tu equipo, el cliente inicia sesión con su correo y contraseña.",
            "Pantalla de login: desde el acceso principal, el usuario completa correo y contraseña y luego entra al lanzador de aplicaciones.",
            "Lanzador de módulos: el sistema muestra las aplicaciones visibles según el rol del usuario. No todos verán las mismas opciones.",
            "Selección de sucursal: cuando la empresa trabaja con varias sedes, es importante confirmar en qué sucursal se está operando antes de cargar inventario, vender o revisar reportes.",
        ],
    )
    add_callout(
        doc,
        "Punto de control",
        [
            "Si un usuario no ve un módulo, primero revisa su rol y permisos.",
            "La primera persona que debe validar el sistema es el administrador del cliente, no el cajero.",
        ],
        fill=AMBER_SOFT,
    )

    add_banner_heading(doc, "3. Configuración inicial obligatoria", "Antes de operar, el cliente debe dejar listos los parámetros que afectan tickets, cobros y acceso.")
    add_two_col_table(
        doc,
        ("Módulo", "Qué se configura aquí"),
        [
            ("Ajustes Generales", "Nombre comercial, NIT, teléfono, dirección, moneda, símbolo, mensaje de ticket, impuestos y facturación electrónica FEL si aplica."),
            ("Sucursales", "Nombre, código, dirección, teléfono y sucursal principal. Cada sede debe existir antes de repartir operación."),
            ("Equipo y permisos", "Usuarios, roles, sucursal base y acceso por áreas. Esto controla quién puede vender, configurar y revisar reportes."),
            ("Bancos y Tesorería", "Cuentas bancarias para cobros con tarjeta o transferencia y control financiero posterior."),
        ],
    )
    add_bullets(
        doc,
        [
            "Ajustes Generales: debe completarlo el administrador. Aquí también se habilitan o deshabilitan efectivo, tarjeta, transferencia y crédito.",
            "FEL: solo debe activarse cuando el cliente ya tenga listo su certificador, NIT emisor y credenciales.",
            "Usuarios: conviene crear al menos un administrador, un supervisor y los cajeros que usarán el POS.",
            "Permisos: el acceso a inventario, reportes, contabilidad y configuración no debe darse a todos por defecto.",
        ],
    )
    add_callout(
        doc,
        "Importante para cobros",
        [
            "Si el negocio aceptará tarjeta o transferencia, debe configurar al menos una cuenta bancaria activa.",
            "En el POS, los pagos con tarjeta y transferencia exigen referencia y banco de destino.",
            "Las ventas a crédito solo se pueden registrar si el cliente está creado en el directorio.",
        ],
    )

    add_banner_heading(doc, "4. Carga del catálogo base", "Sin catálogos correctos, el equipo no podrá operar con orden ni trazabilidad.")
    add_body(doc, "La recomendación es cargar primero productos, después clientes y finalmente proveedores.")
    if INVENTORY_IMAGE.exists():
        add_image(doc, INVENTORY_IMAGE, 6.45, "Vista ilustrativa del control de inventario y catálogo.")
    add_two_col_table(
        doc,
        ("Catálogo", "Buenas prácticas de carga inicial"),
        [
            ("Inventario", "Registrar categorías, productos, SKU, costo, precio, stock inicial, stock mínimo, unidad de medida, exenciones y variantes cuando existan."),
            ("Clientes", "Crear clientes frecuentes, clientes corporativos y clientes con límite de crédito antes de vender al crédito o generar cotizaciones."),
            ("Proveedores", "Cargar razón social, NIT, contacto, teléfono, correo y dirección para poder recibir compras e historial logístico."),
        ],
    )
    add_bullets(
        doc,
        [
            "Inventario: el sistema permite nuevos productos, combos y carga masiva por Excel.",
            "Clientes: desde este módulo también se controlan saldos, límites autorizados y abonos.",
            "Proveedores: conviene registrar primero a los abastecedores activos para que las compras queden bien asociadas.",
        ],
    )

    add_banner_heading(doc, "5. Dejar lista la operación diaria", "Pasos mínimos que el cliente debe completar antes de empezar a vender.")
    add_numbered(
        doc,
        [
            "Confirmar sucursal activa: verificar que el selector de sucursal apunta a la tienda correcta.",
            "Revisar stock inicial: validar que los productos visibles tengan existencia y precios correctos.",
            "Abrir caja: el POS no permite vender si la caja está cerrada. El usuario debe registrar el fondo inicial de efectivo.",
            "Hacer una venta de prueba: agregar un producto, elegir cliente si aplica y completar el cobro.",
            "Revisar el ticket y el historial: confirmar que la venta aparece en ventas recientes, reportes y dashboard.",
        ],
    )
    if POS_IMAGE.exists():
        add_image(doc, POS_IMAGE, 6.45, "Vista ilustrativa del flujo operativo del POS.")
    add_callout(
        doc,
        "Reglas del POS que conviene explicar al cliente",
        [
            "Cotizaciones: para guardar una cotización se debe seleccionar un cliente registrado.",
            "Crédito: una venta a crédito requiere cliente creado y con política interna del negocio.",
            "Pago mixto: el sistema permite combinar métodos de pago.",
            "Egresos: los retiros o gastos hechos desde caja afectan el cierre del turno.",
            "Cierre: al finalizar la jornada, el cajero registra el efectivo final y bloquea la terminal.",
        ],
    )

    add_banner_heading(doc, "6. Compras, inventario y control", "Cómo reabastecer y mantener consistencia en existencias.")
    add_body(
        doc,
        "El flujo de abastecimiento inicia en Proveedores y continúa en Ingresos. Cada recepción de inventario debe quedar asociada a un proveedor y a una referencia documental cuando exista.",
    )
    add_bullets(
        doc,
        [
            "Nueva recepción: seleccionar proveedor, ingresar referencia y agregar productos con cantidad y costo.",
            "Histórico de ingresos: sirve para auditar compras ya registradas.",
            "Solo bajo stock: desde Inventario se pueden filtrar productos que requieren reabastecimiento.",
            "Traslados: cuando haya varias sucursales, los movimientos deben hacerse desde el módulo de traslados para no desordenar existencias.",
        ],
    )
    add_callout(
        doc,
        "Práctica recomendada",
        [
            "No modificar costos o precios sin un responsable definido.",
            "Usar SKU y categorías consistentes desde el primer día.",
            "Registrar compras reales antes de vender productos nuevos si se quiere trazabilidad completa.",
        ],
        fill=GREEN_SOFT,
    )

    add_banner_heading(doc, "7. Monitoreo y seguimiento", "Dónde revisar si la operación está saliendo bien.")
    add_body(
        doc,
        "Una vez iniciada la operación, el cliente debe acostumbrarse a revisar dashboard, historial de ventas y reportes de caja al cierre de cada jornada.",
    )
    if ANALYTICS_IMAGE.exists():
        add_image(doc, ANALYTICS_IMAGE, 6.45, "Vista ilustrativa del dashboard y los indicadores del negocio.")
    add_two_col_table(
        doc,
        ("Pantalla", "Para qué sirve"),
        [
            ("Dashboard", "Ver ventas del día, cantidad de transacciones, productos activos y alertas de bajo stock."),
            ("Ventas", "Consultar historial, filtrar por fecha, estado o canal, revisar cotizaciones y reimprimir tickets."),
            ("Reportes", "Exportar ventas a CSV o PDF y revisar información de caja cuando el turno está abierto."),
            ("Clientes", "Controlar saldos pendientes y registrar abonos de cuentas por cobrar."),
        ],
    )
    add_bullets(
        doc,
        [
            "Cierre diario: revisar que el efectivo final coincida con lo esperado y documentar diferencias.",
            "Seguimiento comercial: usar el historial de ventas y las cotizaciones para retomar clientes.",
            "Análisis básico: confirmar ventas, métodos de pago y alertas de inventario todos los días.",
        ],
    )

    add_banner_heading(doc, "8. Buenas prácticas para el cliente", "Recomendaciones simples que reducen errores desde la primera semana.")
    add_bullets(
        doc,
        [
            "Trabajar siempre con usuarios individuales y no compartir contraseñas.",
            "Verificar sucursal activa antes de vender, comprar o consultar reportes.",
            "Capacitar primero al administrador y luego al personal operativo.",
            "Cargar clientes antes de ofrecer crédito o cotizaciones formales.",
            "Registrar egresos y cierres el mismo día para no perder trazabilidad de caja.",
            "Usar productos con SKU, costo, precio y stock mínimo completos.",
            "Revisar reportes al final del turno, no varios días después.",
        ],
    )
    add_callout(
        doc,
        "Si algo no cuadra",
        [
            "Revisa permisos del usuario.",
            "Confirma sucursal activa.",
            "Valida que existan productos, clientes o cuentas bancarias necesarias para ese flujo.",
            "Si el problema afecta ventas o cierres, detener la operación y escalar de inmediato a soporte.",
        ],
        fill=AMBER_SOFT,
    )

    add_banner_heading(doc, "9. Checklist final de puesta en marcha", "Usar esta tabla como validación antes de declarar listo al cliente.")
    add_checklist_table(
        doc,
        [
            ("Empresa", "Datos fiscales y comerciales completos.", "Tickets y configuración muestran información correcta."),
            ("Sucursales", "Todas las sedes activas y con código definido.", "Los usuarios pueden operar en la sucursal adecuada."),
            ("Usuarios", "Equipo creado con roles y permisos correctos.", "Cada persona ve solo los módulos que necesita."),
            ("Inventario", "Productos clave cargados con precio, costo y stock.", "El POS puede vender sin improvisaciones."),
            ("Clientes", "Clientes frecuentes o corporativos ya registrados.", "Se pueden emitir cotizaciones, ventas a crédito y abonos."),
            ("Proveedores", "Proveedores principales activos.", "Se pueden registrar ingresos a bodega."),
            ("Bancos", "Cuentas disponibles para tarjeta y transferencia.", "El cobro no se bloquea por falta de banco destino."),
            ("Caja", "Apertura, venta de prueba y cierre controlado.", "La operación diaria queda validada de extremo a extremo."),
        ],
    )

    add_banner_heading(doc, "10. Siguientes pasos recomendados", "Cierre sugerido de la implementación después del arranque inicial.")
    add_callout(
        doc,
        "Recomendación para la primera semana",
        [
            "Definir a una persona responsable por configuración, otra por inventario y otra por caja.",
            "Hacer acompañamiento diario del cierre de caja y del dashboard durante los primeros 5 días de operación.",
            "Registrar cualquier duda operativa por módulo para ajustar la capacitación del cliente.",
        ],
        fill=GREEN_SOFT,
    )
    add_bullets(
        doc,
        [
            "Soporte interno: conservar este manual como referencia base para nuevas altas de personal.",
            "Revisión semanal: validar ventas, egresos, inventario y saldos de clientes al menos una vez por semana.",
            "Escalación: si se detecta un error que afecte cobros, stock o cierres, detener el flujo y escalar con evidencia.",
        ],
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_manual()
